import csv
import io
import mimetypes
import os
import re
import statistics
from dataclasses import dataclass


MAX_CONTEXT_CHARS = 28_000
MAX_ROWS = 80
MAX_TEXT_CHARS = 150_000

STOP_WORDS = {
    "a",
    "about",
    "an",
    "and",
    "any",
    "are",
    "as",
    "at",
    "based",
    "be",
    "can",
    "could",
    "document",
    "for",
    "from",
    "give",
    "how",
    "i",
    "in",
    "is",
    "it",
    "me",
    "of",
    "on",
    "or",
    "please",
    "policy",
    "question",
    "share",
    "show",
    "tell",
    "that",
    "the",
    "this",
    "to",
    "uploaded",
    "what",
    "when",
    "where",
    "which",
    "who",
    "with",
    "you",
}


@dataclass
class LoadedFile:
    filename: str
    kind: str
    summary: str
    content: str


def load_file(filename, raw_bytes):
    ext = os.path.splitext(filename.lower())[1]
    if ext == ".csv":
        return _load_csv(filename, raw_bytes)
    if ext in {".xlsx", ".xlsm", ".xltx", ".xltm"}:
        return _load_excel(filename, raw_bytes)
    if ext == ".pdf":
        return _load_pdf(filename, raw_bytes)
    if ext == ".docx":
        return _load_docx(filename, raw_bytes)
    return _load_generic(filename, raw_bytes)


def build_context(files, question=None):
    if question:
        text = _build_relevant_context(files, question)
    else:
        text = "\n\n---\n\n".join(
            f"FILE: {item.filename}\nTYPE: {item.kind}\nSUMMARY:\n{item.summary}\nCONTENT:\n{item.content}"
            for item in files
        )
    if len(text) > MAX_CONTEXT_CHARS:
        return text[:MAX_CONTEXT_CHARS] + "\n\n[Context truncated because the uploaded files are large.]"
    return text


def _build_relevant_context(files, question):
    chunks = []
    for item in files:
        chunks.append(
            {
                "order": len(chunks),
                "score": 1,
                "filename": item.filename,
                "kind": item.kind,
                "title": "File summary",
                "text": item.summary,
            }
        )
        for title, text in _split_content(item.content):
            chunks.append(
                {
                    "order": len(chunks),
                    "score": _score_chunk(question, title, text),
                    "filename": item.filename,
                    "kind": item.kind,
                    "title": title,
                    "text": text,
                }
            )

    ranked = sorted(chunks, key=lambda chunk: chunk["score"], reverse=True)
    if _is_display_request(question):
        ranked = _include_neighbor_chunks(ranked, chunks)

    selected = []
    total_chars = 0
    for chunk in ranked:
        if chunk["score"] <= 0 and selected:
            continue
        formatted = (
            f"FILE: {chunk['filename']}\n"
            f"TYPE: {chunk['kind']}\n"
            f"LOCATION: {chunk['title']}\n"
            f"CONTENT:\n{chunk['text'].strip()}"
        )
        if total_chars + len(formatted) > MAX_CONTEXT_CHARS and selected:
            continue
        selected.append(formatted)
        total_chars += len(formatted)
        if total_chars >= MAX_CONTEXT_CHARS:
            break

    return "\n\n---\n\n".join(selected)


def _include_neighbor_chunks(ranked, chunks):
    expanded = []
    seen_orders = set()
    top_orders = [chunk["order"] for chunk in ranked[:5] if chunk["score"] > 0]
    for order in top_orders:
        for neighbor_order in (order - 1, order, order + 1):
            if 0 <= neighbor_order < len(chunks) and neighbor_order not in seen_orders:
                neighbor = dict(chunks[neighbor_order])
                neighbor["score"] = max(neighbor["score"], 0.5)
                expanded.append(neighbor)
                seen_orders.add(neighbor_order)

    for chunk in ranked:
        if chunk["order"] not in seen_orders:
            expanded.append(chunk)
            seen_orders.add(chunk["order"])

    return expanded


def _split_content(content):
    if not content.strip():
        return []

    page_matches = list(re.finditer(r"(?m)^Page \d+:", content))
    if page_matches:
        chunks = []
        for index, match in enumerate(page_matches):
            start = match.start()
            end = page_matches[index + 1].start() if index + 1 < len(page_matches) else len(content)
            page_text = content[start:end].strip()
            title = page_text.splitlines()[0].strip(":")
            chunks.extend(_split_large_chunk(title, page_text))
        return chunks

    sheet_matches = list(re.finditer(r"(?m)^SHEET: .+$", content))
    if sheet_matches:
        chunks = []
        for index, match in enumerate(sheet_matches):
            start = match.start()
            end = sheet_matches[index + 1].start() if index + 1 < len(sheet_matches) else len(content)
            sheet_text = content[start:end].strip()
            title = sheet_text.splitlines()[0]
            chunks.extend(_split_large_chunk(title, sheet_text))
        return chunks

    return _split_large_chunk("Document text", content)


def _split_large_chunk(title, text, max_chars=4_500):
    cleaned = text.strip()
    if len(cleaned) <= max_chars:
        return [(title, cleaned)]

    parts = []
    paragraphs = re.split(r"\n\s*\n", cleaned)
    current = []
    current_len = 0
    part_number = 1
    for paragraph in paragraphs:
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        if current and current_len + len(paragraph) > max_chars:
            parts.append((f"{title}, part {part_number}", "\n\n".join(current)))
            part_number += 1
            current = []
            current_len = 0
        current.append(paragraph)
        current_len += len(paragraph)
    if current:
        parts.append((f"{title}, part {part_number}", "\n\n".join(current)))
    return parts


def _score_chunk(question, title, text):
    query_terms = _terms(question)
    if not query_terms:
        return 0

    haystack = f"{title}\n{text}".lower()
    compact_haystack = _compact_text(haystack)
    score = 0
    for term in query_terms:
        count = haystack.count(term)
        if count:
            score += min(count, 8) * (3 if term in title.lower() else 1)

    question_phrase = re.sub(r"\s+", " ", question.lower()).strip(" ?.")
    if len(question_phrase) > 8 and question_phrase in haystack:
        score += 12

    compact_question = _compact_text(question)
    if len(compact_question) > 8 and compact_question in compact_haystack:
        score += 20

    for phrase in _important_phrases(question):
        if phrase in haystack or _compact_text(phrase) in compact_haystack:
            score += 10

    return score


def _important_phrases(text):
    cleaned = re.sub(r"[^a-zA-Z0-9\s]", " ", text.lower())
    words = [word for word in cleaned.split() if word not in STOP_WORDS]
    phrases = []
    for size in (4, 3, 2):
        for index in range(0, len(words) - size + 1):
            phrase = " ".join(words[index:index + size])
            if len(phrase) > 7:
                phrases.append(phrase)
    return phrases


def _is_display_request(question):
    text = question.lower()
    display_words = ("share", "show", "display", "provide", "give", "list", "explain")
    document_words = ("policy", "subpolicy", "sub-policy", "form", "annexure", "procedure", "process")
    return any(word in text for word in display_words) and any(word in text for word in document_words)


def _compact_text(text):
    return re.sub(r"[^a-z0-9]", "", text.lower())


def _terms(text):
    words = re.findall(r"[a-zA-Z][a-zA-Z0-9'-]*", text.lower())
    terms = []
    for word in words:
        word = word.strip("'-")
        if len(word) < 3 or word in STOP_WORDS:
            continue
        terms.append(word)
        if word.endswith("ies") and len(word) > 4:
            terms.append(word[:-3] + "y")
        elif word.endswith("es") and len(word) > 4:
            terms.append(word[:-2])
        elif word.endswith("s") and len(word) > 3:
            terms.append(word[:-1])
    return list(dict.fromkeys(terms))


def _load_csv(filename, raw_bytes):
    text = _decode_bytes(raw_bytes)
    rows = []
    for index, row in enumerate(csv.reader(io.StringIO(text))):
        rows.append(row)
        if index >= MAX_ROWS:
            break
    return _table_result(filename, "CSV spreadsheet", rows, total_hint=_line_count(text))


def _load_excel(filename, raw_bytes):
    try:
        import openpyxl
    except ImportError as exc:
        raise ValueError("Excel support needs openpyxl. Run: pip install -r requirements.txt") from exc

    workbook = openpyxl.load_workbook(io.BytesIO(raw_bytes), read_only=True, data_only=True)
    chunks = []
    summaries = []
    for sheet_name in workbook.sheetnames:
        worksheet = workbook[sheet_name]
        rows = []
        for index, row in enumerate(worksheet.iter_rows(values_only=True)):
            rows.append(["" if value is None else str(value) for value in row])
            if index >= MAX_ROWS:
                break
        loaded = _table_result(
            f"{filename} / {sheet_name}",
            "Excel worksheet",
            rows,
            total_hint=worksheet.max_row,
            include_name=False,
        )
        summaries.append(f"{sheet_name}: {worksheet.max_row} rows x {worksheet.max_column} columns")
        chunks.append(f"SHEET: {sheet_name}\n{loaded.content}")
    return LoadedFile(filename, "Excel workbook", "\n".join(summaries), "\n\n".join(chunks))


def _load_pdf(filename, raw_bytes):
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ValueError("PDF support needs pypdf. Run: pip install -r requirements.txt") from exc

    reader = PdfReader(io.BytesIO(raw_bytes))
    pages = []
    for index, page in enumerate(reader.pages):
        pages.append(f"Page {index + 1}:\n{(page.extract_text() or '').strip()}")
    content = "\n\n".join(pages)
    return LoadedFile(
        filename,
        "PDF document",
        f"{len(reader.pages)} pages. Extracted {len(content)} text characters.",
        _limit_text(content),
    )


def _load_docx(filename, raw_bytes):
    try:
        import docx
    except ImportError as exc:
        raise ValueError("DOCX support needs python-docx. Run: pip install -r requirements.txt") from exc

    document = docx.Document(io.BytesIO(raw_bytes))
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    tables = []
    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows[:MAX_ROWS]:
            rows.append(" | ".join(cell.text.strip() for cell in row.cells))
        tables.append(f"Table {table_index}:\n" + "\n".join(rows))
    content = "\n\n".join(paragraphs + tables)
    return LoadedFile(
        filename,
        "Word document",
        f"{len(paragraphs)} paragraphs and {len(document.tables)} tables.",
        _limit_text(content),
    )


def _load_generic(filename, raw_bytes):
    mime_type, _ = mimetypes.guess_type(filename)
    text = _decode_bytes(raw_bytes)
    if _printable_ratio(text) < 0.65:
        content = (
            "This appears to be a binary or unsupported file. "
            f"Size: {len(raw_bytes)} bytes. MIME guess: {mime_type or 'unknown'}."
        )
        return LoadedFile(filename, "unsupported/binary file", content, content)
    return LoadedFile(
        filename,
        mime_type or "text-like file",
        f"Decoded as text with {len(text)} characters.",
        _limit_text(text),
    )


def _table_result(filename, kind, rows, total_hint=None, include_name=True):
    if not rows:
        return LoadedFile(filename, kind, "No rows found.", "")

    headers = rows[0]
    data_rows = rows[1:]
    summary_lines = []
    if include_name:
        summary_lines.append(f"{filename}: {total_hint or len(rows)} rows detected.")
    summary_lines.append(f"Columns: {', '.join(_clean_cell(h) for h in headers)}")
    summary_lines.extend(_numeric_summary(headers, data_rows))
    return LoadedFile(filename, kind, "\n".join(summary_lines), _format_table(rows))


def _numeric_summary(headers, rows):
    summaries = []
    for col_index, header in enumerate(headers[:30]):
        values = []
        for row in rows:
            if col_index >= len(row):
                continue
            try:
                values.append(float(str(row[col_index]).replace(",", "").strip()))
            except ValueError:
                continue
        if values:
            label = _clean_cell(header) or f"Column {col_index + 1}"
            summaries.append(
                f"{label}: count={len(values)}, min={min(values):.4g}, "
                f"max={max(values):.4g}, mean={statistics.fmean(values):.4g}"
            )
    return summaries


def _format_table(rows):
    return "\n".join(" | ".join(_clean_cell(cell) for cell in row) for row in rows[: MAX_ROWS + 1])


def _decode_bytes(raw_bytes):
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return raw_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw_bytes.decode("utf-8", errors="replace")


def _limit_text(text):
    cleaned = text.strip()
    if len(cleaned) > MAX_TEXT_CHARS:
        return cleaned[:MAX_TEXT_CHARS] + "\n\n[Text truncated because the file is large.]"
    return cleaned


def _line_count(text):
    return text.count("\n") + (1 if text else 0)


def _clean_cell(value):
    return str(value).replace("\n", " ").replace("\r", " ").strip()


def _printable_ratio(text):
    if not text:
        return 1
    printable = sum(1 for char in text if char.isprintable() or char.isspace())
    return printable / len(text)